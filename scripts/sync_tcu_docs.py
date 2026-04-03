from __future__ import annotations

import csv
import shutil
import tempfile
from collections import Counter, OrderedDict
from datetime import datetime
from itertools import count
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

try:
    import pythoncom
    from win32com.client import DispatchEx
except Exception:  # pragma: no cover - Windows COM is optional at import time
    pythoncom = None
    DispatchEx = None

ROOT = Path(__file__).resolve().parents[1]
MEETING_DIR = ROOT / "MEETING ADJUSTMENTS"
IMAGE_DIR = ROOT / "images"
UPDATED_LABEL = f"Updated {datetime.now().strftime('%B %d, %Y').replace(' 0', ' ')}"

CSV_PATH = MEETING_DIR / "TCU_Stats_Claim_Inventory.csv"
FACT_CHECK_DOCX_PATH = MEETING_DIR / "TCU_Stats_Fact_Check.docx"
FACT_CHECK_PDF_PATH = MEETING_DIR / "TCU_Stats_Fact_Check.pdf"
DESKTOP_FACT_CHECK_PDF_PATH = ROOT.parent / "tcu-stats-fact-check.pdf"

TEI_LOGO_SOURCE = IMAGE_DIR / "TEILogo.webp"
TCU_LOGO_SOURCE = IMAGE_DIR / "tcu-logo-black.png"

REQUIRED_COLUMNS = [
    "claim_id", "section", "bucket", "claim", "source", "verification", "note",
    "support_type", "public_status", "display_rank", "headline_rank", "subgroup",
    "source_tier", "public_claim", "review_note", "date_sensitive", "last_reviewed",
]

VALID_SUPPORT_TYPES = {"direct", "derived", "interpretation"}
VALID_PUBLIC_STATUS = {"approved", "internal-only", "remove"}
VALID_SOURCE_TIERS = {"primary", "institutional", "derived", "interpretive"}
VALID_DATE_SENSITIVE = {"yes", "no"}

SECTION_PURPOSES = {
    "OVERVIEW": "Quick-glance headline claims used to orient cold reviewers before they move deeper into the evidence base.",
    "Carbon Pools": "Defines the scale of soil carbon, where it sits in the global system, and why losses from soil matter to the atmosphere.",
    "Soil Microbiome": "Supports the biological mechanism layer behind soil carbon storage, fungal transport, and resilience claims.",
    "Carbon Debt - 12,000 Years of Loss": "Documents the historic soil-carbon loss, its CO2 equivalent, and the drawdown opportunity framing.",
    "Rebuilding Carbon": "Supports the rates, practices, and long-term trial evidence used to describe the rebuild pathway.",
    "Profitability": "Validates the core claim that regenerative systems can improve margins while reducing climate-linked input risk.",
    "Case Studies": "Anchors the argument in named farm examples, platform-scale market signals, and applied business outcomes.",
    "Farm Financial Stress": "Shows how climate volatility and debt pressure are already destabilizing conventional farm economics.",
    "Climate & Yields": "Documents observed yield losses, climate-linked crop pressure, and the resilience case for regenerative management.",
    "Soil & Water": "Connects soil organic matter to water holding capacity, infiltration, and climate adaptation on the farm.",
    "Aquifer Crisis": "Supports the groundwater depletion and recharge claims that frame water stress as a climate-linked systems problem.",
    "Central Valley": "Grounds the water crisis in a regional case study with infrastructure, food-system, and land-value consequences.",
    "SGMA & Fallowing": "Documents the compliance and land-retirement pressure used to frame regenerative water retention as a strategic alternative.",
    "Global Scarcity": "Extends the water argument from California to the global food-water-climate collision.",
}

SUPPORT_LABELS = {"direct": "Direct", "derived": "Derived", "interpretation": "Interpretation"}
SOURCE_TIER_LABELS = {
    "primary": "Primary Paper",
    "institutional": "Institutional Dataset",
    "derived": "Derived Calculation",
    "interpretive": "Interpretive Synthesis",
}

GREEN = RGBColor(47, 82, 2)
GREEN_SOFT = "E9F2D9"
GREEN_PALE = "F5F9EE"
GREEN_TINT = "F8FBF3"
EARTH_SOFT = "F4EEE7"
EARTH_PALE = "FAF5EE"
RED_SOFT = "F8E4DE"
TEXT = RGBColor(29, 33, 25)
TEXT_MUTED = RGBColor(91, 98, 84)
ACCENT_RED = RGBColor(192, 57, 43)
ACCENT_EARTH = RGBColor(139, 94, 60)
ACCENT_BLUE = RGBColor(61, 122, 165)
BLUE_SOFT = "EEF4FB"

SUPPORT_FILL = {"direct": "EDF6E1", "derived": "F8F2E6", "interpretation": "FDEBDF"}
SUPPORT_TEXT = {"direct": GREEN, "derived": ACCENT_EARTH, "interpretation": ACCENT_RED}
SOURCE_TIER_FILL = {
    "primary": "EEF4E4", "institutional": "EEF4FB", "derived": "F8F2E6", "interpretive": "FDEBDF",
}
SOURCE_TIER_TEXT = {
    "primary": GREEN, "institutional": ACCENT_BLUE, "derived": ACCENT_EARTH, "interpretive": ACCENT_RED,
}
STATUS_FILL = {"approved": GREEN_SOFT, "internal-only": "FFF6D8", "remove": "FCE7E2"}
STATUS_TEXT = {"approved": GREEN, "internal-only": ACCENT_EARTH, "remove": ACCENT_RED}
BOOKMARK_IDS = count(1)

SECTION_MAP_GROUPS = [
    {
        "title": "Overview",
        "description": "Start here for the top-level orientation and strongest cross-site proof points.",
        "sections": ["OVERVIEW"],
        "fill": GREEN_SOFT,
        "accent": GREEN,
    },
    {
        "title": "Carbon & Soil",
        "description": "Carbon stocks, soil biology, long-term loss, and rebuild pathways.",
        "sections": ["Carbon Pools", "Soil Microbiome", "Carbon Debt - 12,000 Years of Loss", "Rebuilding Carbon"],
        "fill": GREEN_PALE,
        "accent": GREEN,
    },
    {
        "title": "Economics",
        "description": "Farm-business performance, case studies, and climate-linked financial stress.",
        "sections": ["Profitability", "Case Studies", "Farm Financial Stress", "Climate & Yields"],
        "fill": EARTH_SOFT,
        "accent": ACCENT_EARTH,
    },
    {
        "title": "Water",
        "description": "Water retention, aquifer decline, Central Valley impacts, and scarcity.",
        "sections": ["Soil & Water", "Aquifer Crisis", "Central Valley", "SGMA & Fallowing", "Global Scarcity"],
        "fill": BLUE_SOFT,
        "accent": ACCENT_BLUE,
    },
]

def normalize_text(value: str) -> str:
    replacements = {
        "Ã¢â‚¬â€œ": "-", "Ã¢â‚¬â€": "-", "Ã¢â‚¬": '"', "Ã¢â‚¬â„¢": "'", "Ã‚Â°C": "C",
        "Ãƒâ€”": "x", "Ã¢â‚¬Â¢": "*", "Ã¢â‚¬Ë˜": "*", "Ã¢â€°Ë†": "~", "Ã¢â€°Â¥": ">=",
        "Ã¢â€°Â¤": "<=", "1?C": "1C",
    }
    text = value or ""
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.strip()

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def set_fixed_layout(table, widths: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.xpath("./w:tblLayout")
    if tbl_layout:
        tbl_layout[0].set(qn("w:type"), "fixed")
    else:
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)
    for col_idx, width in enumerate(widths):
        if col_idx < len(table.columns):
            table.columns[col_idx].width = Inches(width)
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = Inches(width)

def clear_cell(cell) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_borders(
    table,
    *,
    top: dict[str, str] | None = None,
    bottom: dict[str, str] | None = None,
    left: dict[str, str] | None = None,
    right: dict[str, str] | None = None,
    inside_h: dict[str, str] | None = None,
    inside_v: dict[str, str] | None = None,
) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name, spec in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
        "insideH": inside_h,
        "insideV": inside_v,
    }.items():
        if spec is None:
            continue
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        for key, value in spec.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_borders(
    cell,
    *,
    top: dict[str, str] | None = None,
    bottom: dict[str, str] | None = None,
    left: dict[str, str] | None = None,
    right: dict[str, str] | None = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for name, spec in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        if spec is None:
            continue
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        for key, value in spec.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, *, top=70, bottom=70, left=90, right=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

def style_paragraph(paragraph, *, align=None, before=0, after=0) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)

def add_text_run(paragraph, text: str, *, font="Aptos", size=9.0, bold=False, color=None, italic=False) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_cell_text(cell, text: str, *, font="Aptos", size=9.0, bold=False, color=None, italic=False, align=None, before=0, after=0) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, align=align, before=before, after=after)
    add_text_run(paragraph, text, font=font, size=size, bold=bold, color=color, italic=italic)

def append_cell_line(cell, text: str, *, font="Aptos", size=9.0, bold=False, color=None, italic=False, align=None, before=0, after=0) -> None:
    paragraph = cell.add_paragraph()
    style_paragraph(paragraph, align=align, before=before, after=after)
    add_text_run(paragraph, text, font=font, size=size, bold=bold, color=color, italic=italic)


def bookmark_name(section_name: str) -> str:
    safe = "".join(char if char.isalnum() else "_" for char in section_name)
    return f"sec_{safe}"[:40]


def add_bookmark_to_run(run, name: str) -> None:
    bookmark_id = str(next(BOOKMARK_IDS))
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    run._r.addprevious(start)
    run._r.addnext(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, *, color: str = "2F5202", bold: bool = True) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def rgb_to_hex(color: RGBColor) -> str:
    return "".join(f"{channel:02X}" for channel in color)


def slim_note(text: str, *, limit: int = 78) -> str:
    first_sentence = text.split(". ")[0].strip()
    candidate = first_sentence if first_sentence else text.strip()
    if len(candidate) <= limit:
        return candidate
    truncated = candidate[: limit - 1].rstrip()
    if " " in truncated:
        truncated = truncated.rsplit(" ", 1)[0]
    return f"{truncated}..."


def add_evidence_stack(cell, row: dict[str, str], *, center: bool = True) -> None:
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    add_cell_text(
        cell,
        SUPPORT_LABELS[row["support_type"]],
        font="Aptos Display",
        size=8.8,
        bold=True,
        color=SUPPORT_TEXT[row["support_type"]],
        align=align,
        after=0,
    )
    append_cell_line(
        cell,
        SOURCE_TIER_LABELS[row["source_tier"]],
        font="Aptos Display",
        size=7.6,
        bold=True,
        color=SOURCE_TIER_TEXT[row["source_tier"]],
        align=align,
        after=0,
    )
    if row["date_sensitive"] == "yes":
        append_cell_line(
            cell,
            "Refresh",
            font="Aptos Display",
            size=7.2,
            bold=True,
            color=ACCENT_RED,
            align=align,
            after=0,
        )

def set_page_layout(document: Document, *, landscape: bool = True, margins: float = 0.45) -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(margins)
    section.bottom_margin = Inches(margins)
    section.left_margin = Inches(margins)
    section.right_margin = Inches(margins)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)
    section.different_first_page_header_footer = True

def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.0)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Aptos Display"
            style.font.color.rgb = TEXT

def load_rows() -> list[dict[str, str]]:
    with CSV_PATH.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if not rows:
        raise ValueError("Claim inventory is empty.")
    missing = [column for column in REQUIRED_COLUMNS if column not in rows[0]]
    if missing:
        raise ValueError(f"Claim inventory is missing required columns: {', '.join(missing)}")
    for row in rows:
        for key, value in row.items():
            row[key] = normalize_text(value)
    validate_rows(rows)
    return rows

def validate_rows(rows: list[dict[str, str]]) -> None:
    ordered_sections: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
    for idx, row in enumerate(rows, start=2):
        support_type = row["support_type"].lower()
        public_status = row["public_status"].lower()
        source_tier = row["source_tier"].lower()
        date_sensitive = row["date_sensitive"].lower()

        if support_type not in VALID_SUPPORT_TYPES:
            raise ValueError(f"Row {idx} has invalid support_type '{row['support_type']}'.")
        if public_status not in VALID_PUBLIC_STATUS:
            raise ValueError(f"Row {idx} has invalid public_status '{row['public_status']}'.")
        if source_tier not in VALID_SOURCE_TIERS:
            raise ValueError(f"Row {idx} has invalid source_tier '{row['source_tier']}'.")
        if date_sensitive not in VALID_DATE_SENSITIVE:
            raise ValueError(f"Row {idx} has invalid date_sensitive '{row['date_sensitive']}'.")
        if not row["public_claim"]:
            raise ValueError(f"Row {idx} is missing public_claim.")
        if not row["review_note"]:
            raise ValueError(f"Row {idx} is missing review_note.")
        if not row["last_reviewed"]:
            raise ValueError(f"Row {idx} is missing last_reviewed.")

        try:
            row["_display_rank"] = int(row["display_rank"])
        except ValueError as exc:
            raise ValueError(f"Row {idx} has non-integer display_rank '{row['display_rank']}'.") from exc

        headline_rank = row["headline_rank"]
        if headline_rank:
            try:
                row["_headline_rank"] = int(headline_rank)
            except ValueError as exc:
                raise ValueError(f"Row {idx} has non-integer headline_rank '{headline_rank}'.") from exc
        else:
            row["_headline_rank"] = None

        ordered_sections.setdefault(row["section"], []).append(row)

    for section_name, section_rows in ordered_sections.items():
        display_ranks = sorted(row["_display_rank"] for row in section_rows)
        expected = list(range(1, len(section_rows) + 1))
        if display_ranks != expected:
            raise ValueError(f"Section '{section_name}' must have contiguous display_rank values {expected}; found {display_ranks}.")

        headline_rows = [row for row in section_rows if row["_headline_rank"] is not None]
        if len(headline_rows) > 3:
            raise ValueError(f"Section '{section_name}' has more than 3 headline_rank rows.")
        if any(row["public_status"] != "approved" for row in headline_rows):
            raise ValueError(f"Section '{section_name}' has headline_rank rows that are not approved.")

        headline_ranks = sorted(row["_headline_rank"] for row in headline_rows)
        expected_headlines = list(range(1, len(headline_rows) + 1))
        if headline_ranks != expected_headlines:
            raise ValueError(f"Section '{section_name}' must have contiguous headline_rank values {expected_headlines}; found {headline_ranks}.")

def ordered_section_groups(rows: list[dict[str, str]]) -> OrderedDict[str, list[dict[str, str]]]:
    groups: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
    for row in rows:
        groups.setdefault(row["section"], []).append(row)
    for section_name in groups:
        groups[section_name].sort(key=lambda item: item["_display_rank"])
    return groups

def ensure_logo_assets() -> tuple[Path, Path, Path]:
    temp_dir = Path(tempfile.mkdtemp(prefix="tcu_factcheck_"))
    tei_png = temp_dir / "tei-logo.png"
    with Image.open(TEI_LOGO_SOURCE) as img:
        img.save(tei_png, format="PNG")
    return tei_png, TCU_LOGO_SOURCE, temp_dir

def add_paragraph(document: Document, text: str, *, size: float = 9.0, bold: bool = False, color: RGBColor | None = None, align: int | None = None, before: float = 0, after: float = 4, uppercase: bool = False) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, align=align, before=before, after=after)
    add_text_run(paragraph, text.upper() if uppercase else text, size=size, bold=bold, color=color)

def support_mix_text(rows: list[dict[str, str]]) -> str:
    counts = Counter(row["support_type"] for row in rows)
    pieces = []
    for key in ("direct", "derived", "interpretation"):
        count = counts.get(key, 0)
        if count:
            pieces.append(f"{count} {key}")
    return " | ".join(pieces)

def add_cover(document: Document, approved_rows: list[dict[str, str]], appendix_rows: list[dict[str, str]], ordered_sections: list[str], tei_logo: Path, tcu_logo: Path) -> None:
    top = document.add_table(rows=1, cols=3)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_layout(top, [1.9, 5.4, 1.9])

    left = top.cell(0, 0)
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    left.paragraphs[0].add_run().add_picture(str(tei_logo), width=Inches(1.55))

    center = top.cell(0, 1)
    center.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker = center.paragraphs[0]
    add_text_run(kicker, "TCU Stats Review Packet", font="Aptos Display", size=9.5, bold=True, color=GREEN)
    title = center.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FACT-CHECK REGISTER")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = TEXT
    subtitle = center.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(subtitle, "Approved public-facing claim register for TCU Stats.", size=10.5, color=TEXT_MUTED)
    subtitle_2 = center.add_paragraph()
    subtitle_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(subtitle_2, "Designed for faster cold review, clearer evidence signaling, and cleaner source verification.", size=9.8, color=TEXT_MUTED)

    right = top.cell(0, 2)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraphs[0].add_run().add_picture(str(tcu_logo), width=Inches(1.35))

    # Push the metrics block lower on the cover so page one opens with more
    # breathing room before the approved-claims summary.
    add_paragraph(document, "", after=28)

    metrics = document.add_table(rows=1, cols=3)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_layout(metrics, [3.0, 3.0, 3.0])
    metric_values = [
        ("Approved Claims", str(len(approved_rows)), GREEN_SOFT, GREEN),
        ("Sections", str(len(ordered_sections)), EARTH_SOFT, GREEN),
        ("Appendix Rows", str(len(appendix_rows)), RED_SOFT, ACCENT_RED),
    ]
    for idx, (label, value, fill, text_color) in enumerate(metric_values):
        cell = metrics.cell(0, idx)
        set_cell_shading(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(22)
        run.font.color.rgb = text_color
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(paragraph, label, size=9, color=TEXT_MUTED)

    add_paragraph(document, "This register is the verification layer behind TCU Stats. The main body contains only approved rows for public-facing review; internal-only and removed claims are preserved later in a separate appendix.", size=10.3, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, before=8)

    document.add_page_break()

def add_running_header_footer(document: Document, tei_logo: Path, tcu_logo: Path) -> None:
    section = document.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(9.7))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_layout(header_table, [1.5, 6.1, 1.5])
    header_table.cell(0, 0).paragraphs[0].add_run().add_picture(str(tei_logo), width=Inches(1.15))

    center = header_table.cell(0, 1).paragraphs[0]
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = center.add_run("TCU Stats Fact-Check Register")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN

    right = header_table.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run().add_picture(str(tcu_logo), width=Inches(0.95))

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(3)
    run = footer_paragraph.add_run(f"The Edison Institute | The Carbon Underground | Fact-check review pack | {UPDATED_LABEL}")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_MUTED

def add_methodology_page(document: Document, rows: list[dict[str, str]], ordered_sections: list[str]) -> None:
    add_paragraph(document, "How To Read This Register", size=18, bold=True, color=TEXT, after=8)
    add_paragraph(document, "Review flow is intentionally layered: scan the ranked claims first, use trust signals to judge evidence strength, then move into the full register when you need source-level detail.", size=10.2, color=TEXT_MUTED, after=10)

    intro = document.add_table(rows=1, cols=3)
    intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_layout(intro, [3.0, 3.0, 3.0])
    blocks = [
        ("Review Path", "Headline Claims -> Full Register -> Appendix. The main body only includes approved rows."),
        ("Evidence Labels", "Support Type tells you how the claim was built. Source Tier tells you what kind of evidence it rests on."),
        ("Fast Scan", "Each section opens with three ranked headline claims before the full table so a reviewer can orient quickly."),
    ]
    for idx, (title, body) in enumerate(blocks):
        cell = intro.cell(0, idx)
        set_cell_shading(cell, GREEN_PALE if idx != 1 else EARTH_SOFT)
        add_cell_text(cell, title, font="Aptos Display", size=11.2, bold=True, color=GREEN)
        append_cell_line(cell, body, size=8.8, color=TEXT_MUTED)

    add_paragraph(document, "Support Types", size=13.5, bold=True, color=GREEN, after=5, before=10)
    support = document.add_table(rows=1, cols=2)
    support.style = "Table Grid"
    support.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(support, [2.0, 7.0])
    for idx, header in enumerate(["Display", "Meaning"]):
        cell = support.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, GREEN_SOFT)
    set_repeat_header(support.rows[0])
    support_rows = [
        ("Direct", "Directly stated in the cited source or dataset."),
        ("Derived", "Explicitly calculated from cited source values; the math lives in verification."),
        ("Interpretation", "Editorial synthesis grounded in cited evidence rather than a direct source sentence."),
    ]
    for display, meaning in support_rows:
        cells = support.add_row().cells
        cells[0].text = display
        cells[1].text = meaning
        set_cell_shading(cells[0], SUPPORT_FILL[display.lower()])
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = SUPPORT_TEXT[display.lower()]
    style_table_font(support, 8.1)

    add_paragraph(document, "Source Tiers", size=13.5, bold=True, color=GREEN, after=5, before=8)
    tiers = document.add_table(rows=1, cols=2)
    tiers.style = "Table Grid"
    tiers.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(tiers, [2.0, 7.0])
    for idx, header in enumerate(["Display", "Meaning"]):
        cell = tiers.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, GREEN_SOFT)
    set_repeat_header(tiers.rows[0])
    tier_rows = [
        ("Primary Paper", "Peer-reviewed or primary research source."),
        ("Institutional Dataset", "Government, institutional, program, or market report source."),
        ("Derived Calculation", "Reviewer-facing rollup built from approved source values."),
        ("Interpretive Synthesis", "Framing layer used to connect evidence into reviewer messaging."),
    ]
    for display, meaning in tier_rows:
        cells = tiers.add_row().cells
        cells[0].text = display
        cells[1].text = meaning
        tier_key = next(key for key, value in SOURCE_TIER_LABELS.items() if value == display)
        set_cell_shading(cells[0], SOURCE_TIER_FILL[tier_key])
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = SOURCE_TIER_TEXT[tier_key]
    style_table_font(tiers, 8.1)

    add_paragraph(document, "Refresh Flag", size=13.5, bold=True, color=GREEN, after=5, before=8)
    refresh = document.add_table(rows=1, cols=2)
    refresh.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(refresh, [1.7, 7.3])
    set_cell_shading(refresh.cell(0, 0), RED_SOFT)
    add_cell_text(refresh.cell(0, 0), "Refresh", font="Aptos Display", size=10.5, bold=True, color=ACCENT_RED, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(refresh.cell(0, 1), GREEN_TINT)
    add_cell_text(refresh.cell(0, 1), "Time-sensitive claim. Recheck before future publication, reuse, or date-sensitive review cycles.", size=8.6, color=TEXT_MUTED)

    add_paragraph(document, "Status Boundary", size=13.5, bold=True, color=GREEN, after=6, before=10)
    status = document.add_table(rows=1, cols=3)
    status.style = "Table Grid"
    status.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(status, [1.25, 4.0, 3.75])
    for idx, header in enumerate(["Status", "Meaning", "Where It Appears"]):
        cell = status.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, GREEN_SOFT)
    set_repeat_header(status.rows[0])
    status_rows = [
        ("APPROVED", "Cleared for public-facing review and site alignment.", "Main body"),
        ("INTERNAL-ONLY", "Useful provenance or context, but not strong enough for public-facing review.", "Appendix"),
        ("REMOVE", "Retained for provenance only; not approved for current public-facing use.", "Appendix"),
    ]
    for label, meaning, where in status_rows:
        cells = status.add_row().cells
        cells[0].text = label
        cells[1].text = meaning
        cells[2].text = where
        set_cell_shading(cells[0], STATUS_FILL[label.lower()])
    style_table_font(status, 8.2)


def add_section_map_page(document: Document, rows: list[dict[str, str]], ordered_sections: list[str]) -> None:
    document.add_page_break()
    add_paragraph(document, "Section Map", size=22, bold=True, color=TEXT, after=4)
    add_paragraph(document, "Clickable contents guide. Start with Overview, then move left to right across Carbon & Soil, Economics, and Water.", size=8.9, color=TEXT_MUTED, after=12)

    counts = Counter(row["section"] for row in rows if row["public_status"] == "approved")
    overview = SECTION_MAP_GROUPS[0]
    overview_table = document.add_table(rows=1, cols=1)
    overview_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(overview_table, [9.35])
    set_table_borders(
        overview_table,
        top={"val": "nil"},
        bottom={"val": "nil"},
        left={"val": "nil"},
        right={"val": "nil"},
        inside_h={"val": "nil"},
        inside_v={"val": "nil"},
    )
    overview_cell = overview_table.cell(0, 0)
    clear_cell(overview_cell)
    set_cell_margins(overview_cell, top=40, bottom=90, left=0, right=0)
    set_cell_borders(overview_cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "D9E4C7"})
    add_cell_text(overview_cell, overview["title"], font="Aptos Display", size=12.8, bold=True, color=overview["accent"], after=1)
    append_cell_line(overview_cell, overview["description"], size=8.1, color=TEXT_MUTED, after=6)

    overview_links = overview_cell.add_table(rows=1, cols=2)
    overview_links.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(overview_links, [7.85, 1.5])
    set_table_borders(
        overview_links,
        top={"val": "nil"},
        bottom={"val": "nil"},
        left={"val": "nil"},
        right={"val": "nil"},
        inside_h={"val": "nil"},
        inside_v={"val": "nil"},
    )
    link_cell = overview_links.cell(0, 0)
    count_cell = overview_links.cell(0, 1)
    clear_cell(link_cell)
    clear_cell(count_cell)
    set_cell_margins(link_cell, top=35, bottom=35, left=0, right=0)
    set_cell_margins(count_cell, top=35, bottom=35, left=0, right=0)
    link_paragraph = link_cell.paragraphs[0]
    style_paragraph(link_paragraph, after=0)
    add_internal_hyperlink(link_paragraph, "OVERVIEW", bookmark_name("OVERVIEW"), color=rgb_to_hex(overview["accent"]))
    count_paragraph = count_cell.paragraphs[0]
    style_paragraph(count_paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_text_run(count_paragraph, str(counts.get("OVERVIEW", 0)), font="Aptos Display", size=10.2, bold=True, color=overview["accent"])

    pillar_grid = document.add_table(rows=1, cols=3)
    pillar_grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(pillar_grid, [3.1, 3.1, 3.1])
    set_table_borders(
        pillar_grid,
        top={"val": "nil"},
        bottom={"val": "nil"},
        left={"val": "nil"},
        right={"val": "nil"},
        inside_h={"val": "nil"},
        inside_v={"val": "nil"},
    )

    for idx, group in enumerate(SECTION_MAP_GROUPS[1:]):
        cell = pillar_grid.cell(0, idx)
        clear_cell(cell)
        set_cell_margins(cell, top=90, bottom=0, left=0, right=0)
        title = cell.paragraphs[0]
        style_paragraph(title, after=1)
        add_text_run(title, group["title"], font="Aptos Display", size=12.1, bold=True, color=group["accent"])
        description = cell.add_paragraph()
        style_paragraph(description, after=7)
        add_text_run(description, group["description"], size=7.7, color=TEXT_MUTED)

        links = cell.add_table(rows=len(group["sections"]), cols=2)
        links.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_fixed_layout(links, [2.55, 0.42])
        set_table_borders(
            links,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
            inside_h={"val": "nil"},
            inside_v={"val": "nil"},
        )
        row_idx = 0
        for section_name in group["sections"]:
            if section_name not in ordered_sections:
                continue
            left_cell = links.cell(row_idx, 0)
            right_cell = links.cell(row_idx, 1)
            clear_cell(left_cell)
            clear_cell(right_cell)
            set_cell_margins(left_cell, top=28, bottom=28, left=0, right=0)
            set_cell_margins(right_cell, top=28, bottom=28, left=0, right=0)
            line_spec = {"val": "single", "sz": "5", "space": "0", "color": "E7ECE1"}
            set_cell_borders(left_cell, bottom=line_spec)
            set_cell_borders(right_cell, bottom=line_spec)
            link = left_cell.paragraphs[0]
            style_paragraph(link, after=0)
            add_internal_hyperlink(link, section_name, bookmark_name(section_name), color=rgb_to_hex(group["accent"]))
            count_line = right_cell.paragraphs[0]
            style_paragraph(count_line, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
            add_text_run(count_line, str(counts.get(section_name, 0)), font="Aptos Display", size=8.3, bold=True, color=group["accent"])
            row_idx += 1

    document.add_page_break()

def section_summary_strip(document: Document, section_name: str, approved_rows: list[dict[str, str]]) -> None:
    title = document.add_paragraph()
    style_paragraph(title, before=2, after=6)
    title_run = title.add_run(section_name.upper())
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = GREEN
    add_bookmark_to_run(title_run, bookmark_name(section_name))

    summary = document.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(summary, [2.2, 7.15])
    set_table_borders(
        summary,
        top={"val": "nil"},
        bottom={"val": "single", "sz": "8", "space": "0", "color": "D9E4C7"},
        left={"val": "nil"},
        right={"val": "nil"},
        inside_h={"val": "nil"},
        inside_v={"val": "nil"},
    )
    left, right = summary.rows[0].cells
    clear_cell(left)
    clear_cell(right)
    set_cell_margins(left, top=20, bottom=85, left=0, right=70)
    set_cell_margins(right, top=28, bottom=85, left=70, right=0)
    add_cell_text(left, str(len(approved_rows)), font="Aptos Display", size=23, bold=True, color=TEXT, after=0)
    append_cell_line(left, "Approved claims", font="Aptos Display", size=9.2, bold=True, color=GREEN, after=0)
    append_cell_line(left, support_mix_text(approved_rows), size=7.8, color=TEXT_MUTED)
    add_cell_text(right, "Purpose", font="Aptos Display", size=8.6, bold=True, color=GREEN, after=1)
    append_cell_line(right, SECTION_PURPOSES.get(section_name, "Section summary used to orient reviewers before row-level verification."), size=8.7, color=TEXT_MUTED)

def headline_rows(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    picks = [row for row in rows if row["public_status"] == "approved" and row["_headline_rank"] is not None]
    picks.sort(key=lambda item: item["_headline_rank"])
    return picks

def add_headline_claims(document: Document, rows: list[dict[str, str]]) -> None:
    picks = headline_rows(rows)
    if not picks:
        return
    add_paragraph(document, "Headline Claims", size=13.5, bold=True, color=GREEN, after=3, before=10)
    add_paragraph(document, "Start here. These are the fastest claims to scan before dropping into the full register.", size=8.3, color=TEXT_MUTED, after=8)

    for row in picks:
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_fixed_layout(table, [7.35, 1.95])
        set_table_borders(
            table,
            top={"val": "nil"},
            bottom={"val": "single", "sz": "8", "space": "0", "color": "DEE6D2"},
            left={"val": "nil"},
            right={"val": "nil"},
            inside_h={"val": "nil"},
            inside_v={"val": "nil"},
        )
        left, right = table.rows[0].cells
        clear_cell(left)
        clear_cell(right)
        set_cell_margins(left, top=40, bottom=55, left=0, right=70)
        set_cell_margins(right, top=52, bottom=55, left=70, right=0)
        add_cell_text(left, f"H{row['_headline_rank']}", font="Aptos Display", size=8.0, bold=True, color=GREEN, after=0)
        if row["subgroup"]:
            append_cell_line(left, row["subgroup"].upper(), font="Aptos Display", size=7.1, bold=True, color=TEXT_MUTED, after=0)
        append_cell_line(left, row["public_claim"], font="Aptos Display", size=12.1, bold=True, color=TEXT, after=2)
        append_cell_line(left, f"{row['source']} | {row['verification']}", size=7.8, color=TEXT_MUTED)
        add_evidence_stack(right, row)

def style_table_font(table, size: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    if not run.font.name:
                        run.font.name = "Aptos"
                    if run.font.size is None:
                        run.font.size = Pt(size)

def add_subgroup_row(table, label: str, fill: str = GREEN_PALE) -> None:
    cells = table.add_row().cells
    merged = cells[0]
    for cell in cells[1:]:
        merged = merged.merge(cell)
    set_cell_shading(merged, fill)
    set_cell_margins(merged, top=24, bottom=24, left=45, right=45)
    set_cell_borders(merged, bottom={"val": "single", "sz": "5", "space": "0", "color": "E6ECE0"})
    add_cell_text(merged, label.upper(), font="Aptos Display", size=7.8, bold=True, color=GREEN)

def add_main_table(document: Document, rows: list[dict[str, str]]) -> None:
    approved_rows = [row for row in rows if row["public_status"] == "approved"]
    add_paragraph(document, "Full Register", size=12.5, bold=True, color=GREEN, after=4, before=8)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(table, [0.55, 3.35, 2.2, 1.0, 2.25])
    set_table_borders(
        table,
        top={"val": "nil"},
        bottom={"val": "nil"},
        left={"val": "nil"},
        right={"val": "nil"},
        inside_h={"val": "nil"},
        inside_v={"val": "nil"},
    )
    headers = ["ID", "Claim", "Source", "Evidence", "Review"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        add_cell_text(cell, header, font="Aptos Display", size=8.2, bold=True, color=GREEN)
        set_cell_shading(cell, GREEN_TINT)
        set_cell_margins(cell, top=24, bottom=24, left=40, right=40)
        set_cell_borders(cell, bottom={"val": "single", "sz": "8", "space": "0", "color": "CBD8B8"})
    set_repeat_header(table.rows[0])

    current_subgroup = None
    for row_idx, row in enumerate(approved_rows):
        subgroup = row["subgroup"] or None
        if subgroup != current_subgroup and subgroup:
            add_subgroup_row(table, subgroup)
        current_subgroup = subgroup

        cells = table.add_row().cells
        row_fill = GREEN_TINT if row_idx % 2 == 0 else "FFFFFF"
        for cell in cells:
            set_cell_shading(cell, row_fill)
            set_cell_margins(cell, top=26, bottom=34, left=45, right=45)
            set_cell_borders(cell, bottom={"val": "single", "sz": "5", "space": "0", "color": "E6ECE0"})

        add_cell_text(cells[0], row["claim_id"], font="Aptos Display", size=7.7, bold=False, color=TEXT, after=0)
        add_cell_text(cells[1], row["public_claim"], size=8.5, bold=True, color=TEXT, after=1)
        if row["subgroup"]:
            append_cell_line(cells[1], row["subgroup"].upper(), font="Aptos Display", size=6.9, bold=True, color=TEXT_MUTED)
        add_cell_text(cells[2], row["source"], size=8.0, color=TEXT, after=1)
        append_cell_line(cells[2], row["verification"], size=7.1, color=TEXT_MUTED)
        add_evidence_stack(cells[3], row)
        add_cell_text(cells[4], slim_note(row["review_note"]), size=7.5, color=TEXT_MUTED)

    style_table_font(table, 7.3)

def add_appendix_divider(document: Document, appendix_rows: list[dict[str, str]]) -> None:
    document.add_page_break()
    add_paragraph(document, "Appendix - Internal Provenance", size=18, bold=True, color=TEXT, after=8)
    add_paragraph(document, "This appendix preserves internal-only and removed rows so provenance is not lost. These pages are not part of the approved cold-review flow.", size=10.2, color=TEXT_MUTED, after=10)

    counts = Counter(row["public_status"] for row in appendix_rows)
    cards = document.add_table(rows=1, cols=3)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_layout(cards, [3.0, 3.0, 3.0])
    card_copy = [
        ("Internal-Only", str(counts.get("internal-only", 0)), "Context kept for internal review, but not cleared for public-facing use.", "FFF6D8", ACCENT_EARTH),
        ("Removed", str(counts.get("remove", 0)), "Rows retained for provenance only and intentionally excluded from current public review.", "FCE7E2", ACCENT_RED),
        ("Main Body Rule", "Approved only", "If a row is not approved, it should appear here rather than in the main section register.", GREEN_TINT, GREEN),
    ]
    for idx, (title, value, body, fill, color) in enumerate(card_copy):
        cell = cards.cell(0, idx)
        set_cell_shading(cell, fill)
        add_cell_text(cell, title, font="Aptos Display", size=11.2, bold=True, color=color)
        append_cell_line(cell, value, font="Aptos Display", size=16, bold=True, color=TEXT)
        append_cell_line(cell, body, size=8.5, color=TEXT_MUTED)

def add_appendix(document: Document, appendix_groups: OrderedDict[str, list[dict[str, str]]]) -> None:
    appendix_rows = [row for rows in appendix_groups.values() for row in rows]
    add_appendix_divider(document, appendix_rows)

    for section_name, rows in appendix_groups.items():
        add_paragraph(document, section_name.upper(), size=12.5, bold=True, color=GREEN, before=8, after=4)
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_fixed_layout(table, [0.52, 0.92, 0.82, 0.98, 2.6, 1.35, 1.76])
        headers = ["ID", "Status", "Support", "Tier", "Claim", "Source", "Notes"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, GREEN_SOFT)
        set_repeat_header(table.rows[0])

        current_subgroup = None
        for row in rows:
            subgroup = row["subgroup"] or None
            if subgroup != current_subgroup and subgroup:
                add_subgroup_row(table, subgroup, fill=EARTH_SOFT)
            current_subgroup = subgroup

            cells = table.add_row().cells
            values = [
                row["claim_id"], row["public_status"].upper(), SUPPORT_LABELS[row["support_type"]],
                SOURCE_TIER_LABELS[row["source_tier"]], row["claim"], row["source"], row["note"],
            ]
            for idx, value in enumerate(values):
                cells[idx].text = value

            set_cell_shading(cells[1], STATUS_FILL[row["public_status"]])
            set_cell_shading(cells[2], SUPPORT_FILL[row["support_type"]])
            set_cell_shading(cells[3], SOURCE_TIER_FILL[row["source_tier"]])
            for idx in (1, 2, 3):
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cells[idx].paragraphs[0].runs:
                    run.bold = True
            for run in cells[1].paragraphs[0].runs:
                run.font.color.rgb = STATUS_TEXT[row["public_status"]]
            for run in cells[2].paragraphs[0].runs:
                run.font.color.rgb = SUPPORT_TEXT[row["support_type"]]
            for run in cells[3].paragraphs[0].runs:
                run.font.color.rgb = SOURCE_TIER_TEXT[row["source_tier"]]

        style_table_font(table, 7.3)

def build_fact_check(rows: list[dict[str, str]]) -> None:
    tei_logo, tcu_logo, temp_dir = ensure_logo_assets()
    try:
        document = Document()
        set_page_layout(document, landscape=True, margins=0.45)
        style_document(document)

        grouped = ordered_section_groups(rows)
        ordered_sections = list(grouped.keys())
        approved_groups: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
        appendix_groups: OrderedDict[str, list[dict[str, str]]] = OrderedDict()

        for section_name, section_rows in grouped.items():
            approved_rows = [row for row in section_rows if row["public_status"] == "approved"]
            appendix_rows = [row for row in section_rows if row["public_status"] != "approved"]
            if approved_rows:
                approved_groups[section_name] = approved_rows
            if appendix_rows:
                appendix_groups[section_name] = appendix_rows

        approved_rows = [row for row in rows if row["public_status"] == "approved"]
        appendix_rows = [row for row in rows if row["public_status"] != "approved"]

        add_cover(document, approved_rows, appendix_rows, ordered_sections, tei_logo, tcu_logo)
        add_running_header_footer(document, tei_logo, tcu_logo)
        add_methodology_page(document, rows, ordered_sections)
        add_section_map_page(document, rows, ordered_sections)

        first_section = True
        for section_name, section_rows in approved_groups.items():
            if not first_section:
                document.add_page_break()
            section_summary_strip(document, section_name, section_rows)
            add_headline_claims(document, section_rows)
            add_main_table(document, section_rows)
            first_section = False

        if appendix_groups:
            add_appendix(document, appendix_groups)

        document.save(FACT_CHECK_DOCX_PATH)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

def export_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    if DispatchEx is None or pythoncom is None:
        raise RuntimeError("Word COM export is unavailable on this machine.")

    pythoncom.CoInitialize()
    word = DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    temp_copy = None
    try:
        temp_copy = Path(tempfile.mkdtemp(prefix="tcu_factcheck_export_")) / docx_path.name
        shutil.copy2(docx_path, temp_copy)
        doc = word.Documents.Open(str(temp_copy), ReadOnly=True, AddToRecentFiles=False, Visible=False, ConfirmConversions=False)
        doc.ExportAsFixedFormat(str(pdf_path), 17)
        doc.Close(False)
    finally:
        word.Quit()
        pythoncom.CoUninitialize()
        if temp_copy is not None:
            shutil.rmtree(temp_copy.parent, ignore_errors=True)

def main() -> None:
    rows = load_rows()
    build_fact_check(rows)
    export_docx_to_pdf(FACT_CHECK_DOCX_PATH, FACT_CHECK_PDF_PATH)
    shutil.copy2(FACT_CHECK_PDF_PATH, DESKTOP_FACT_CHECK_PDF_PATH)

if __name__ == "__main__":
    main()
